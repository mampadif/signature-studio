import base64
import io
import streamlit as st
import numpy as np
from dataclasses import dataclass
from typing import Tuple, Optional, List
from PIL import Image, ImageDraw, ImageFilter, ImageFont, ImageOps, PngImagePlugin

# =========================================================
# 1. IMPORTS & COMPATIBILITY
# =========================================================
try:
    import stripe
    STRIPE_AVAILABLE = True
except ImportError:
    STRIPE_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from google import genai
    HAS_GENAI = True
except ImportError:
    HAS_GENAI = False

# =========================================================
# 2. CONFIGURATION & STATE
# =========================================================
@dataclass
class AppConfig:
    api_key: str = st.secrets.get("GEMINI_API_KEY", "")
    ai_model: str = st.secrets.get("GEMINI_MODEL", "gemini-3.1-flash-image-preview")
    app_name: str = st.secrets.get("APP_NAME", "Signature Studio Pro")
    max_calls: int = int(st.secrets.get("MAX_AI_CALLS_PER_SESSION", 3))
    stripe_sk: str = st.secrets.get("STRIPE_SECRET_KEY", "")
    stripe_price_id: str = st.secrets.get("STRIPE_PRICE_ID_SIGNATURE", "")
    paypal_url: str = st.secrets.get("PAYPAL_PAYMENT_URL", "")
    app_url: str = st.secrets.get("APP_URL", "")
    unlock_code: str = st.secrets.get("UNLOCK_CODE", "1234")
    price: str = st.secrets.get("PRICE_DISPLAY", "$3.99")

CONFIG = AppConfig()

# Modern Design System Colors
COLORS = {
    "primary": "#6366F1",       # Indigo-500
    "primary_hover": "#4F46E5", # Indigo-600
    "primary_light": "#EEF2FF", # Indigo-50
    "secondary": "#0F172A",     # Slate-900
    "accent": "#06B6D4",        # Cyan-500
    "success": "#10B981",       # Emerald-500
    "warning": "#F59E0B",       # Amber-500
    "danger": "#EF4444",        # Red-500
    "bg": "#FAFAFA",            # Neutral-50
    "surface": "#FFFFFF",
    "border": "#E2E8F0",
    "text_primary": "#0F172A",
    "text_secondary": "#64748B",
    "text_muted": "#94A3B8",
    "gradient_hero": "linear-gradient(135deg, #6366F1 0%, #8B5CF6 50%, #06B6D4 100%)",
    "gradient_card": "linear-gradient(135deg, #FFFFFF 0%, #F8FAFC 100%)",
}

st.set_page_config(
    page_title=CONFIG.app_name,
    page_icon="🖊️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

if "paid" not in st.session_state:
    st.session_state.paid = False
if "ai_calls_used" not in st.session_state:
    st.session_state.ai_calls_used = 0
if "final_img" not in st.session_state:
    st.session_state.final_img = None
if "processing_step" not in st.session_state:
    st.session_state.processing_step = 0

# =========================================================
# 3. MODERN PREMIUM UI CSS (2025 Design Patterns)
# =========================================================
st.markdown(f"""
<style>
    /* ===== GLOBAL STYLES ===== */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&display=swap');
    
    .stApp {{
        background: {COLORS["bg"]};
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
    }}

    /* ===== ANIMATIONS ===== */
    @keyframes fadeInUp {{
        from {{ opacity: 0; transform: translateY(20px); }}
        to {{ opacity: 1; transform: translateY(0); }}
    }}
    @keyframes pulse {{
        0%, 100% {{ transform: scale(1); }}
        50% {{ transform: scale(1.05); }}
    }}
    @keyframes shimmer {{
        0% {{ background-position: -1000px 0; }}
        100% {{ background-position: 1000px 0; }}
    }}
    @keyframes float {{
        0%, 100% {{ transform: translateY(0px); }}
        50% {{ transform: translateY(-10px); }}
    }}
    
    .animate-in {{
        animation: fadeInUp 0.6s cubic-bezier(0.16, 1, 0.3, 1) forwards;
    }}
    .animate-in-delay-1 {{ animation-delay: 0.1s; opacity: 0; }}
    .animate-in-delay-2 {{ animation-delay: 0.2s; opacity: 0; }}
    .animate-in-delay-3 {{ animation-delay: 0.3s; opacity: 0; }}

    /* ===== HERO SECTION ===== */
    .hero-wrapper {{
        background: {COLORS["gradient_hero"]};
        border-radius: 28px;
        padding: 4rem 2.5rem;
        text-align: center;
        margin-bottom: 2.5rem;
        position: relative;
        overflow: hidden;
        box-shadow: 0 25px 50px -12px rgba(99, 102, 241, 0.25);
    }}
    .hero-wrapper::before {{
        content: '';
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: radial-gradient(circle, rgba(255,255,255,0.1) 0%, transparent 70%);
        animation: float 6s ease-in-out infinite;
    }}
    .hero-title {{
        font-size: 3.8rem;
        font-weight: 900;
        color: white;
        margin-bottom: 0.75rem;
        letter-spacing: -0.03em;
        position: relative;
        z-index: 1;
        line-height: 1.1;
    }}
    .hero-subtitle {{
        font-size: 1.25rem;
        color: rgba(255,255,255,0.9);
        font-weight: 400;
        position: relative;
        z-index: 1;
        max-width: 600px;
        margin: 0 auto 1.5rem;
        line-height: 1.6;
    }}
    .hero-badge {{
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        background: rgba(255,255,255,0.2);
        backdrop-filter: blur(10px);
        padding: 0.5rem 1.25rem;
        border-radius: 999px;
        color: white;
        font-size: 0.9rem;
        font-weight: 600;
        position: relative;
        z-index: 1;
        border: 1px solid rgba(255,255,255,0.3);
        margin-bottom: 1.5rem;
    }}

    /* ===== CARD STYLES ===== */
    .card {{
        background: {COLORS["surface"]};
        border-radius: 20px;
        padding: 2rem;
        border: 1px solid {COLORS["border"]};
        box-shadow: 0 1px 3px rgba(0,0,0,0.04);
        transition: all 0.3s cubic-bezier(0.16, 1, 0.3, 1);
    }}
    .card:hover {{
        border-color: {COLORS["primary"]}20;
        box-shadow: 0 8px 30px rgba(99, 102, 241, 0.08);
        transform: translateY(-2px);
    }}
    .card-accent {{
        border-left: 4px solid {COLORS["primary"]};
    }}

    /* ===== RESULT CONTAINER ===== */
    .result-wrapper {{
        background: {COLORS["surface"]};
        border-radius: 24px;
        border: 2px solid {COLORS["border"]};
        padding: 2.5rem;
        text-align: center;
        position: relative;
        overflow: hidden;
    }}
    .result-wrapper::after {{
        content: '';
        position: absolute;
        inset: 0;
        background-image: 
            linear-gradient(45deg, #f1f5f9 25%, transparent 25%),
            linear-gradient(-45deg, #f1f5f9 25%, transparent 25%),
            linear-gradient(45deg, transparent 75%, #f1f5f9 75%),
            linear-gradient(-45deg, transparent 75%, #f1f5f9 75%);
        background-size: 20px 20px;
        background-position: 0 0, 0 10px, 10px -10px, -10px 0px;
        opacity: 0.3;
        z-index: 0;
    }}
    .result-content {{
        position: relative;
        z-index: 1;
    }}

    /* ===== PAYMENT CARD ===== */
    .payment-card {{
        background: {COLORS["gradient_card"]};
        border: 2px solid {COLORS["primary_light"]};
        border-radius: 24px;
        padding: 2.5rem;
        text-align: center;
        box-shadow: 0 4px 24px rgba(99, 102, 241, 0.06);
    }}
    .payment-card h2 {{
        color: {COLORS["secondary"]};
        font-size: 1.75rem;
        font-weight: 800;
        margin-bottom: 0.5rem;
    }}
    .payment-card .price-tag {{
        font-size: 3rem;
        font-weight: 900;
        color: {COLORS["primary"]};
        margin: 1rem 0;
    }}

    /* ===== BUTTON STYLES ===== */
    .btn-primary {{
        background: {COLORS["primary"]};
        color: white;
        border: none;
        padding: 0.9rem 2rem;
        border-radius: 14px;
        font-weight: 700;
        font-size: 1.05rem;
        cursor: pointer;
        transition: all 0.2s ease;
        box-shadow: 0 4px 14px rgba(99, 102, 241, 0.3);
        text-decoration: none;
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        justify-content: center;
    }}
    .btn-primary:hover {{
        background: {COLORS["primary_hover"]};
        transform: translateY(-2px);
        box-shadow: 0 8px 24px rgba(99, 102, 241, 0.4);
    }}
    .btn-secondary {{
        background: white;
        color: {COLORS["secondary"]};
        border: 2px solid {COLORS["border"]};
        padding: 0.9rem 2rem;
        border-radius: 14px;
        font-weight: 700;
        font-size: 1.05rem;
        cursor: pointer;
        transition: all 0.2s ease;
        text-decoration: none;
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        justify-content: center;
    }}
    .btn-secondary:hover {{
        border-color: {COLORS["primary"]};
        color: {COLORS["primary"]};
        transform: translateY(-2px);
    }}
    .btn-success {{
        background: {COLORS["success"]};
        color: white;
        border: none;
        padding: 0.9rem 2rem;
        border-radius: 14px;
        font-weight: 700;
        font-size: 1.05rem;
        cursor: pointer;
        transition: all 0.2s ease;
        box-shadow: 0 4px 14px rgba(16, 185, 129, 0.3);
        text-decoration: none;
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        justify-content: center;
    }}
    .btn-success:hover {{
        box-shadow: 0 8px 24px rgba(16, 185, 129, 0.4);
        transform: translateY(-2px);
    }}

    /* ===== STATUS BADGES ===== */
    .badge {{
        display: inline-flex;
        align-items: center;
        gap: 0.35rem;
        padding: 0.35rem 0.85rem;
        border-radius: 999px;
        font-size: 0.8rem;
        font-weight: 600;
    }}
    .badge-success {{
        background: #D1FAE5;
        color: #065F46;
    }}
    .badge-warning {{
        background: #FEF3C7;
        color: #92400E;
    }}
    .badge-info {{
        background: {COLORS["primary_light"]};
        color: {COLORS["primary"]};
    }}

    /* ===== PROGRESS INDICATOR ===== */
    .progress-steps {{
        display: flex;
        gap: 1rem;
        align-items: center;
        padding: 1rem 0;
    }}
    .progress-step {{
        display: flex;
        align-items: center;
        gap: 0.5rem;
        color: {COLORS["text_muted"]};
        font-size: 0.9rem;
        font-weight: 500;
    }}
    .progress-step.active {{
        color: {COLORS["primary"]};
        font-weight: 700;
    }}
    .progress-step.done {{
        color: {COLORS["success"]};
    }}
    .progress-dot {{
        width: 10px;
        height: 10px;
        border-radius: 50%;
        background: {COLORS["border"]};
    }}
    .progress-dot.active {{
        background: {COLORS["primary"]};
        box-shadow: 0 0 0 4px {COLORS["primary_light"]};
    }}
    .progress-dot.done {{
        background: {COLORS["success"]};
        box-shadow: 0 0 0 4px #D1FAE5;
    }}

    /* ===== SIDEBAR UPGRADE ===== */
    [data-testid="stSidebar"] {{
        background: linear-gradient(180deg, #0F172A 0%, #1E293B 100%);
        border-right: none;
    }}
    [data-testid="stSidebar"] * {{
        color: #E2E8F0 !important;
    }}
    [data-testid="stSidebar"] h1,
    [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] h3 {{
        color: #F1F5F9 !important;
    }}
    [data-testid="stSidebar"] .stSlider label {{
        color: #CBD5E1 !important;
    }}

    /* ===== TOOLTIP STYLE ===== */
    .tooltip {{
        position: relative;
        display: inline-block;
        cursor: help;
        border-bottom: 1px dotted {COLORS["text_muted"]};
    }}

    /* ===== RESPONSIVE ===== */
    @media (max-width: 768px) {{
        .hero-title {{ font-size: 2.2rem; }}
        .hero-wrapper {{ padding: 2rem 1rem; }}
    }}

    /* ===== FILE UPLOADER ===== */
    [data-testid="stFileUploader"] {{
        transition: all 0.3s ease;
    }}
    [data-testid="stFileUploader"]:hover {{
        border-color: {COLORS["primary"]} !important;
    }}

    /* ===== FOOTER ===== */
    .footer {{
        text-align: center;
        padding: 2rem;
        color: {COLORS["text_muted"]};
        font-size: 0.85rem;
        border-top: 1px solid {COLORS["border"]};
        margin-top: 3rem;
    }}
    .footer a {{
        color: {COLORS["primary"]};
        text-decoration: none;
    }}
</style>
""", unsafe_allow_html=True)

# =========================================================
# 4. EXTRACTION ENGINE (Preserved from original)
# =========================================================
def validate_quality(img: Image.Image) -> Tuple[bool, str]:
    """Ensures input meets technical standards for extraction."""
    arr = np.array(img.convert("RGB"))
    h, w, _ = arr.shape
    if w < 250 or h < 250:
        return False, "Photo resolution too low for pro extraction."
    gray = (0.299*arr[:,:,0] + 0.587*arr[:,:,1] + 0.114*arr[:,:,2]).astype(np.uint8)
    ink_ratio = float((gray < 185).sum()) / (w * h)
    if ink_ratio < 0.0002:
        return False, "No clear signature detected. Try better lighting."
    return True, "Quality Validated"

def process_pipeline(original_img, a_thresh, softness):
    """The master pipeline combining AI unboxing and soft-edge logic."""
    img = ImageOps.exif_transpose(original_img)
    
    if HAS_GENAI and CONFIG.api_key and st.session_state.ai_calls_used < CONFIG.max_calls:
        try:
            client = genai.Client(api_key=CONFIG.api_key)
            response = client.models.generate_content(
                model=CONFIG.ai_model,
                contents=["Extract ONLY the signature ink. High contrast black ink on pure white background.", img]
            )
            for part in response.candidates[0].content.parts:
                if hasattr(part, 'inline_data') or hasattr(part, 'as_image'):
                    img = part.as_image()
                    st.session_state.ai_calls_used += 1
                    break
        except Exception:
            pass

    img = img.convert("RGBA")
    arr = np.array(img)
    brightness = arr[:, :, :3].mean(axis=2)
    
    soft_start = a_thresh - softness
    alpha = np.where(brightness >= a_thresh, 0,
             np.where(brightness <= soft_start, 255,
             ((a_thresh - brightness) / (a_thresh - soft_start) * 255)))
    
    arr[:, :, 3] = alpha.astype(np.uint8)
    ink = alpha > 0
    arr[ink, 0:3] = np.minimum(arr[ink, 0:3], 25)
    
    final = Image.fromarray(arr)
    bbox = final.getchannel("A").getbbox()
    if bbox:
        final = final.crop(bbox)
    
    return final

# =========================================================
# 5. MODERNIZED APP WORKFLOW
# =========================================================

# ---- SIDEBAR (Dark Theme) ----
with st.sidebar:
    st.markdown("""
    <div style="text-align: center; padding: 1rem 0;">
        <svg width="48" height="48" viewBox="0 0 24 24" fill="none" stroke="#6366F1" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
            <path d="M13 2L3 14h9l-1 8 10-12h-9l1-8z"/>
        </svg>
        <h2 style="margin-top: 0.5rem;">Engine Settings</h2>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("---")
    st.markdown("#### 🎚️ Extraction Parameters")
    
    a_thresh = st.slider(
        "Transparency Threshold",
        200, 255, 248,
        help="Higher values remove more white paper background"
    )
    softness = st.slider(
        "Edge Smoothing",
        5, 45, 18,
        help="Higher values create softer edges around strokes"
    )
    
    st.markdown("---")
    st.markdown("#### 📊 Session Stats")
    
    quota_used = st.session_state.ai_calls_used
    quota_total = CONFIG.max_calls
    quota_pct = (quota_used / quota_total * 100) if quota_total > 0 else 0
    
    cols = st.columns([3, 1])
    with cols[0]:
        st.caption(f"AI Quota Used")
    with cols[1]:
        st.markdown(f'<span class="badge badge-info">{quota_used}/{quota_total}</span>', unsafe_allow_html=True)
    
    st.progress(quota_pct / 100, text=f"{int(quota_pct)}%")
    
    st.markdown("---")
    if st.button("🔄 Start Fresh Project", use_container_width=True):
        st.session_state.clear()
        st.session_state.ai_calls_used = 0
        st.session_state.final_img = None
        st.session_state.processing_step = 0
        st.rerun()

# ---- HERO SECTION ----
st.markdown(f"""
<div class="hero-wrapper animate-in">
    <div class="hero-badge animate-in animate-in-delay-1">
        <span>✨</span> AI-Powered Professional Extraction
    </div>
    <h1 class="hero-title animate-in animate-in-delay-1">{CONFIG.app_name}</h1>
    <p class="hero-subtitle animate-in animate-in-delay-2">
        Transform any handwritten signature into a crisp, transparent digital asset 
        ready for Word, PDF, and all your business documents.
    </p>
</div>
""", unsafe_allow_html=True)

# ---- PROGRESS INDICATOR ----
if st.session_state.final_img:
    current_step = 3
elif st.session_state.processing_step > 0:
    current_step = 2
else:
    current_step = 1

steps = [
    ("1", "Upload Photo", current_step >= 1),
    ("2", "Configure & Process", current_step >= 2),
    ("3", "Export & Download", current_step >= 3),
]

st.markdown('<div class="progress-steps">', unsafe_allow_html=True)
for i, (num, label, active) in enumerate(steps):
    step_class = "done" if current_step > (i + 1) else ("active" if current_step == (i + 1) else "")
    dot_class = "done" if current_step > (i + 1) else ("active" if current_step == (i + 1) else "")
    
    st.markdown(f'''
    <div class="progress-step {step_class}">
        <div class="progress-dot {dot_class}"></div>
        <span>{num}. {label}</span>
    </div>
    ''', unsafe_allow_html=True)
    
    if i < len(steps) - 1:
        st.markdown(f'<span style="color: {"#10B981" if current_step > (i+1) else "#E2E8F0"};">→</span>', unsafe_allow_html=True)
st.markdown('</div>', unsafe_allow_html=True)

# ---- MAIN DASHBOARD ----
upload_col, preview_col = st.columns([1, 1], gap="large")

with upload_col:
    st.markdown("""
    <div class="card card-accent animate-in animate-in-delay-1">
        <h3 style="margin-top: 0;">📂 1. Upload Source Photo</h3>
        <p style="color: #64748B; font-size: 0.9rem;">Choose a clear photo of your signature on white paper.</p>
    </div>
    """, unsafe_allow_html=True)
    
    file = st.file_uploader(
        "",
        type=['png', 'jpg', 'jpeg', 'webp'],
        label_visibility="collapsed",
        key="signature_upload"
    )
    
    if file:
        original = Image.open(file)
        
        # Inline preview mini
        with st.container():
            st.markdown("#### 📸 Photo Preview")
            mini = original.copy()
            mini.thumbnail((300, 200))
            st.image(mini, use_container_width=False)
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        if st.button("✨ Generate Pro Signature", type="primary", use_container_width=True):
            is_valid, msg = validate_quality(original)
            if not is_valid:
                st.error(f"⚠️ {msg}")
                st.info("💡 **Tip:** Make sure your signature is on clean white paper with good lighting.")
            else:
                st.session_state.processing_step = 2
                with st.spinner("🔄 AI is analyzing stroke patterns..."):
                    st.session_state.final_img = process_pipeline(original, a_thresh, softness)
                    st.session_state.processing_step = 3
                st.success("✅ Signature extracted successfully!")
                st.balloons()

with preview_col:
    st.markdown("""
    <div class="card card-accent animate-in animate-in-delay-2">
        <h3 style="margin-top: 0;">🎯 2. Live Preview</h3>
        <p style="color: #64748B; font-size: 0.9rem;">Your processed signature appears here in real-time.</p>
    </div>
    """, unsafe_allow_html=True)
    
    if st.session_state.final_img:
        st.markdown('<div class="result-wrapper">', unsafe_allow_html=True)
        st.markdown('<div class="result-content">', unsafe_allow_html=True)
        
        display_img = st.session_state.final_img.copy()
        if not st.session_state.paid:
            draw = ImageDraw.Draw(display_img)
            try:
                font = ImageFont.truetype("DejaVuSans-Bold.ttf", max(16, display_img.width // 18))
            except:
                font = ImageFont.load_default()
            
            watermark_text = "PREVIEW • UNLOCK TO DOWNLOAD"
            bbox = draw.textbbox((0, 0), watermark_text, font=font)
            tw = bbox[2] - bbox[0]
            th = bbox[3] - bbox[1]
            
            # Repeating watermark pattern
            for y in range(0, display_img.height, th + 40):
                for x in range(0, display_img.width, tw + 60):
                    draw.text((x, y), watermark_text, fill=(180, 180, 180, 128), font=font)
        
        st.image(display_img, use_container_width=False)
        
        if not st.session_state.paid:
            st.markdown(
                f'<br><span class="badge badge-warning">🔒 Watermark Active — Unlock for clean export</span>',
                unsafe_allow_html=True
            )
        else:
            st.markdown(
                f'<br><span class="badge badge-success">✅ Verified Purchase — Full access active</span>',
                unsafe_allow_html=True
            )
        
        st.markdown('</div></div>', unsafe_allow_html=True)
    else:
        st.info("""
        👋 **Your signature will appear here after processing.**
        
        1. Upload your photo using the panel on the left
        2. Adjust the extraction settings in the sidebar if needed
        3. Click "Generate Pro Signature"
        """)

# ---- MONETIZATION & EXPORT SECTION ----
if st.session_state.final_img:
    st.markdown("<br>", unsafe_allow_html=True)

    if st.session_state.paid:
        st.markdown("""
        <div class="animate-in">
            <div style="background: linear-gradient(135deg, #D1FAE5, #A7F3D0); 
                        padding: 1.5rem; border-radius: 18px; text-align: center; margin-bottom: 1.5rem;">
                <h2 style="color: #065F46; margin: 0;">✅ Payment Confirmed — Your files are ready</h2>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        dl_col1, dl_col2 = st.columns(2, gap="large")
        
        with dl_col1:
            st.markdown('<div class="card" style="text-align: center;">', unsafe_allow_html=True)
            st.markdown("### 📄 Transparent PNG")
            st.caption("Best for Word, Google Docs, and PDF insertion")
            buf = io.BytesIO()
            st.session_state.final_img.save(buf, format="PNG", optimize=True)
            st.download_button(
                "⬇️ Download PNG",
                buf.getvalue(),
                "signature_studio_pro.png",
                "image/png",
                use_container_width=True,
                type="primary"
            )
            st.markdown('</div>', unsafe_allow_html=True)
        
        with dl_col2:
            if DOCX_AVAILABLE:
                st.markdown('<div class="card" style="text-align: center;">', unsafe_allow_html=True)
                st.markdown("### 📝 Word-Ready DOCX")
                st.caption("Pre-formatted with placement instructions")
                doc = Document()
                doc.add_heading("Signature Asset", level=1)
                doc.add_paragraph("Place this image above your signature line. Choose Layout → In Front of Text.")
                img_s = io.BytesIO()
                st.session_state.final_img.save(img_s, format="PNG")
                doc.add_picture(img_s, width=Inches(2.5))
                doc_buf = io.BytesIO()
                doc.save(doc_buf)
                st.download_button(
                    "⬇️ Download DOCX",
                    doc_buf.getvalue(),
                    "signature_pro.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    type="secondary"
                )
                st.markdown('</div>', unsafe_allow_html=True)
            else:
                st.warning("Install python-docx to enable Word-ready download.")
    else:
        # Payment section with modern design
        st.markdown("""
        <div class="payment-card animate-in">
            <span class="badge badge-info">🔒 Premium Feature</span>
            <h2 style="margin-top: 1rem;">Unlock High-Resolution Export</h2>
            <p style="color: #64748B;">Remove the watermark and access professional file formats</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("<br>", unsafe_allow_html=True)
        pay_col1, pay_col2 = st.columns(2, gap="large")
        
        with pay_col1:
            st.markdown('<div class="card" style="text-align: center;">', unsafe_allow_html=True)
            st.markdown(f"### 💳 Card Payment")
            st.markdown(f"<p class='price-tag'>{CONFIG.price}</p>", unsafe_allow_html=True)
            st.caption("One-time purchase • Instant unlock")
            
            if st.button("💳 Pay via Stripe", use_container_width=True, key="stripe_btn"):
                if STRIPE_AVAILABLE and CONFIG.stripe_sk and CONFIG.stripe_price_id:
                    try:
                        stripe.api_key = CONFIG.stripe_sk
                        sess = stripe.checkout.Session.create(
                            mode="payment",
                            line_items=[{"price": CONFIG.stripe_price_id, "quantity": 1}],
                            success_url=f"{CONFIG.app_url}?paid=1&session_id={{CHECKOUT_SESSION_ID}}",
                            cancel_url=CONFIG.app_url
                        )
                        st.markdown(
                            f'<meta http-equiv="refresh" content="0;URL=\'{sess.url}\'" />',
                            unsafe_allow_html=True
                        )
                    except Exception as e:
                        st.error("Stripe not configured. Please use PayPal or access code.")
                else:
                    st.warning("Stripe not configured. Use PayPal or access code.")
            st.markdown('</div>', unsafe_allow_html=True)
        
        with pay_col2:
            st.markdown('<div class="card" style="text-align: center;">', unsafe_allow_html=True)
            st.markdown(f"### 🔵 PayPal")
            st.markdown(f"<p class='price-tag'>{CONFIG.price}</p>", unsafe_allow_html=True)
            st.caption("Secure payment • Instant unlock")
            
            if CONFIG.paypal_url:
                st.link_button("🔵 Pay via PayPal", CONFIG.paypal_url, use_container_width=True)
            else:
                st.warning("PayPal URL not configured.")
            st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown("<br>", unsafe_allow_html=True)
        with st.expander("🎫 Already purchased? Enter access code", expanded=False):
            code = st.text_input("Enter your access code", type="password", key="unlock_code")
            if st.button("Verify & Unlock", use_container_width=True, key="verify_btn"):
                if code.strip() == CONFIG.unlock_code.strip():
                    st.session_state.paid = True
                    st.success("🎉 Access granted! Your downloads are now available.")
                    st.rerun()
                else:
                    st.error("Invalid code. Please try again.")

# ---- FOOTER ----
st.markdown("""
<div class="footer">
    <p>© 2026 Technoworks Pty Ltd · Signature Studio Pro</p>
    <p style="font-size: 0.8rem; margin-top: 0.25rem;">
        Professional-grade AI signature extraction · 
        <a href="#">Privacy Policy</a> · 
        <a href="#">Terms of Service</a>
    </p>
</div>
""", unsafe_allow_html=True)